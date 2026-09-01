"""Backend tests for /api/pdf/pdf-to-word: header + valid docx bytes."""
import io, os, zipfile, requests, pytest
from docx import Document
from openpyxl import load_workbook

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL').rstrip('/')


def _post(endpoint, files, data=None):
    # Simulate browser CORS request so expose_headers appears
    headers = {'Origin': BASE_URL}
    return requests.post(f'{BASE_URL}/api/pdf/{endpoint}', files=files, data=data or {}, headers=headers, timeout=180)


class TestPdfToWord:
    def test_text_pdf_download_headers_and_validity(self):
        with open('/tmp/text.pdf', 'rb') as f:
            r = _post('pdf-to-word', {'file': ('text.pdf', f, 'application/pdf')})
        assert r.status_code == 200, r.text[:300]
        expose = r.headers.get('access-control-expose-headers', '')
        assert 'Content-Disposition' in expose, f'expose header missing: {expose}'
        cd = r.headers.get('content-disposition', '')
        assert 'text.docx' in cd, f'CD wrong: {cd}'
        data = r.content
        assert zipfile.ZipFile(io.BytesIO(data)).testzip() is None
        doc = Document(io.BytesIO(data))
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'hello world' in text.lower(), text[:200]

    def test_scan_pdf_download_headers_and_validity(self):
        with open('/tmp/scan.pdf', 'rb') as f:
            r = _post('pdf-to-word', {'file': ('scan.pdf', f, 'application/pdf')}, {'lang': 'eng'})
        assert r.status_code == 200, r.text[:300]
        cd = r.headers.get('content-disposition', '')
        assert 'scan.docx' in cd, f'CD wrong: {cd}'
        data = r.content
        assert zipfile.ZipFile(io.BytesIO(data)).testzip() is None
        doc = Document(io.BytesIO(data))
        text = '\n'.join(p.text for p in doc.paragraphs).upper()
        assert 'SCANNED' in text and 'DOCUMENT' in text, text[:300]


class TestPdfToExcel:
    def test_scan_pdf_to_xlsx(self):
        with open('/tmp/scan.pdf', 'rb') as f:
            r = _post('pdf-to-excel', {'file': ('scan.pdf', f, 'application/pdf')})
        assert r.status_code == 200, r.text[:300]
        wb = load_workbook(io.BytesIO(r.content))
        assert len(wb.sheetnames) >= 1
